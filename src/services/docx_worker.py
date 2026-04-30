import asyncio
from pathlib import Path
from typing import AsyncGenerator, Union, List, Optional

from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from core.config import settings
from entities.section import Section
from schemas.docx import HeaderMeta, ChunkMeta
from utils.data_working import get_document_type
from utils.logging import logger

class DocxWorker:
    
    def __init__(self):
        self.chunker = RecursiveCharacterTextSplitter(
            chunk_size=settings.MAX_CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    async def process_document(
        self, 
        file_path: Path
    ) -> AsyncGenerator[Union[HeaderMeta, ChunkMeta], None]:
        logger.info(f"DocxWorker: начало обработки {file_path.name}")
        
        doc_type = get_document_type(file_path)
        source_file = file_path.name
        doc = await asyncio.get_event_loop().run_in_executor(None, lambda: Document(file_path))

        heading_stack: List[str] = []
        text_buffer: List[str] = []

        async def flush_buffer(current_headings: List[str]):
            if not text_buffer:
                return
            
            full_text = "\n".join(text_buffer)
            chunks = self.chunker.split_text(full_text)
            
            for chunk_text in chunks:
                yield ChunkMeta(
                    text=chunk_text,
                    doc_type=doc_type,
                    heading_titles=list(current_headings),
                    source_file=source_file,
                )
            text_buffer.clear()

        for para in doc.paragraphs:
            level = self._get_heading_level(para)
            
            if level is not None:
                async for chunk in flush_buffer(heading_stack):
                    yield chunk

                header_text = para.text.strip()
                if not header_text:
                    continue

                while len(heading_stack) >= level:
                    heading_stack.pop()
                
                parent_titles = list(heading_stack)
                heading_stack.append(header_text)

                yield HeaderMeta(
                    title=header_text,
                    section_level=level,
                    doc_type=doc_type,
                    parent_titles=parent_titles,
                    source_file=source_file,
                )
            
            elif para.text.strip():
                text_buffer.append(para.text.strip())

        async for chunk in flush_buffer(heading_stack):
            yield chunk

        logger.info(f"DocxWorker: завершена обработка {file_path.name}")

    async def process_template_document(
        self, 
        file_path: Path
    ) -> AsyncGenerator[Section, None]:
        logger.info(f"DocxWorker: начало формирования дерева {file_path.name}")

        loop = asyncio.get_event_loop()
        doc = await loop.run_in_executor(None, lambda: Document(file_path))

        current_root: Optional[Section] = None
        stack: List[tuple[int, Section]] = []
        
        min_level = 999
        for para in doc.paragraphs:
            lvl = self._get_heading_level(para)
            if lvl is not None:
                min_level = min(min_level, lvl)

        for para in doc.paragraphs:
            level = self._get_heading_level(para)

            if level is not None:
                new_section = Section(title=para.text.strip(), text="", children=[])

                if level == min_level:
                    if current_root:
                        yield current_root
                    current_root = new_section
                    stack = [(level, current_root)]
                else:
                    while stack and stack[-1][0] >= level:
                        stack.pop()
                    if stack:
                        stack[-1][1].children.append(new_section)
                    stack.append((level, new_section))

            elif para.text.strip():
                if not stack:
                    current_root = Section(title="", text="", children=[])
                    stack = [(0, current_root)]
                
                current_section = stack[-1][1]
                if current_section.text:
                    current_section.text += "\n\n" + para.text
                else:
                    current_section.text = para.text

        if current_root:
            yield current_root

        logger.info("DocxWorker: дерево сформировано")
    
    def _get_heading_level(self, paragraph) -> Optional[int]:
        style_name = paragraph.style.name
        if style_name.startswith('Heading'):
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 1
        return None
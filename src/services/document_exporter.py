import io
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph

from entities.skeleton import Section
from repositories.skeleton_repository import SkeletonRepository

class DocumentExporter:
    
    COLORS = {
        "base": RGBColor(0, 102, 204),    # Синий: Из RAG
        "ai": RGBColor(128, 0, 128),      # Фиолетовый: Придумано ИИ
        "human": RGBColor(204, 0, 0),     # Красный: Заполнить человеку
        "default": RGBColor(0, 0, 0)      # Черный: Обычный текст
    }

    async def build_docx_iteratively(self, chat_id: str, repo: SkeletonRepository) -> io.BytesIO:
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        doc.add_heading('Спецификация / Документ', 0)

        total_sections = await repo.get_total_sections_count(chat_id)

        def process_section(node: Section, level: int = 1):
            doc.add_heading(node.title, level=level)
            if node.text:
                p = doc.add_paragraph()
                self._add_colored_text(p, node.text)
            for child in node.children:
                # Ограничиваем уровень заголовков до 9 (максимум в docx)
                process_section(child, level=min(level + 1, 9))

        for index in range(total_sections):
            section = await repo.get_section_model(chat_id, index)
            if section:
                process_section(section, level=1)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_colored_text(self, paragraph: Paragraph, text: str):
        if not text:
            return
        
        pattern = re.compile(r'\[(base|ai|human)\](.*?)\[/\1\]|([^\[]+)', re.DOTALL)
        
        for match in pattern.finditer(text):
            tag = match.group(1)
            tagged_text = match.group(2)
            untagged_text = match.group(3)
            
            if tag and tagged_text:
                run = paragraph.add_run(tagged_text)
                run.font.color.rgb = DocumentExporter.COLORS[tag]
                if tag == "human":
                    run.font.highlight_color = 7 # Желтый маркер (по желанию)
            elif untagged_text:
                run = paragraph.add_run(untagged_text)
                run.font.color.rgb = DocumentExporter.COLORS["default"]
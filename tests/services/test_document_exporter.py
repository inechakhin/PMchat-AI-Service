import io
import pytest
from unittest.mock import AsyncMock, MagicMock, Mock, patch
from docx.shared import RGBColor

from services.document_exporter import DocumentExporter
from entities.skeleton import Section

class MockRun:
    def __init__(self, text=""):
        self.text = text
        self.font = Mock()
        self.font.color = Mock()
        self.font.color.rgb = None

class MockParagraph:
    def __init__(self):
        self.runs = []
        self.add_run = Mock(side_effect=self._add_run)

    def _add_run(self, text):
        run = MockRun(text)
        self.runs.append(run)
        return run

def extract_headings(mock_add_heading):
    headings = []
    for c in mock_add_heading.call_args_list:
        title = c.args[0]
        if "level" in c.kwargs:
            level = c.kwargs["level"]
        elif len(c.args) > 1:
            level = c.args[1]
        else:
            level = None
        headings.append((title, level))
    return headings

@pytest.fixture
def mock_repo():
    return AsyncMock()

@pytest.fixture
def exporter():
    return DocumentExporter()

def make_section(title, text="", children=None):
    if children is None:
        children = []
    return Section(title=title, text=text, children=children)

@pytest.mark.parametrize("input_text, expected_runs", [
    ("Hello world", [("Hello world", RGBColor(0, 0, 0))]),
    ("[base]base text[/base]", [("base text", RGBColor(0, 0, 0))]),
    ("[ai]AI[/ai]", [("AI", RGBColor(128, 128, 128))]),
    ("[human]Human[/human]", [("Human", RGBColor(204, 0, 0))]),
    ("Start[ai]AI[/ai]End", [
        ("Start", RGBColor(0, 0, 0)),
        ("AI", RGBColor(128, 128, 128)),
        ("End", RGBColor(0, 0, 0)),
    ]),
    ("", []),
])
def test_add_colored_text(exporter, input_text, expected_runs):
    para = MockParagraph()
    exporter._add_colored_text(para, input_text)
    assert len(para.runs) == len(expected_runs)
    for actual_run, (exp_text, exp_color) in zip(para.runs, expected_runs):
        assert actual_run.text == exp_text
        assert actual_run.font.color.rgb == exp_color

@pytest.mark.asyncio
async def test_build_empty(mock_repo, exporter):
    mock_repo.get_total_sections_count.return_value = 0
    with patch("services.document_exporter.Document") as MockDoc:
        mock_doc = MockDoc.return_value
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc.add_heading = MagicMock()
        mock_doc.save = MagicMock()
        result = await exporter.build_docx_iteratively("chat1", mock_repo)

    headings = extract_headings(mock_doc.add_heading)
    assert headings == [("Спецификация / Документ", 0)]
    assert isinstance(result, io.BytesIO)

@pytest.mark.asyncio
async def test_build_with_sections(mock_repo, exporter):
    child = make_section("Child", text="child text")
    sec1 = make_section("First", text="[ai]AI part[/ai] plain")
    sec2 = make_section("Second", text="", children=[child])

    mock_repo.get_total_sections_count.return_value = 2
    mock_repo.get_section_model.side_effect = lambda cid, idx: [sec1, sec2][idx]

    with patch("services.document_exporter.Document") as MockDoc:
        mock_doc = MockDoc.return_value
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc.add_heading = MagicMock()
        mock_para1 = MockParagraph()
        mock_para2 = MockParagraph()
        mock_doc.add_paragraph = MagicMock(side_effect=[mock_para1, mock_para2])
        mock_doc.save = MagicMock()
        result = await exporter.build_docx_iteratively("chat1", mock_repo)

    headings = extract_headings(mock_doc.add_heading)
    assert headings == [
        ("Спецификация / Документ", 0),
        ("First", 1),
        ("Second", 1),
        ("Child", 2),
    ]

    assert len(mock_para1.runs) == 2
    assert mock_para1.runs[0].text == "AI part"
    assert mock_para1.runs[0].font.color.rgb == exporter.COLORS["ai"]
    assert mock_para1.runs[1].text == " plain"
    assert mock_para1.runs[1].font.color.rgb == exporter.COLORS["default"]

    assert len(mock_para2.runs) == 1
    assert mock_para2.runs[0].text == "child text"
    assert mock_para2.runs[0].font.color.rgb == exporter.COLORS["default"]

    mock_doc.save.assert_called_once()
    assert isinstance(result, io.BytesIO)

@pytest.mark.asyncio
async def test_build_section_none(mock_repo, exporter):
    mock_repo.get_total_sections_count.return_value = 2
    mock_repo.get_section_model.side_effect = [None, make_section("Only")]

    with patch("services.document_exporter.Document") as MockDoc:
        mock_doc = MockDoc.return_value
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc.add_heading = MagicMock()
        mock_doc.save = MagicMock()
        await exporter.build_docx_iteratively("chat1", mock_repo)

    headings = extract_headings(mock_doc.add_heading)
    assert headings == [
        ("Спецификация / Документ", 0),
        ("Only", 1),
    ]